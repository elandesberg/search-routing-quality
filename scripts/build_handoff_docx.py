#!/usr/bin/env python3
"""Build the Google-Docs handoff DOCX from the controlled HTML source.

The HTML is the editable source in this package.  This builder does not ask Word
or Google Docs to interpret its CSS, JavaScript, pseudo-elements, or inline SVG.
Instead it maps the source to native Word structures and rasterizes the five
figures with a fixed light palette.
"""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import html as html_stdlib
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree, html


EXPECTED = {
    "h1": 1,
    "h2": 12,
    "h3": 18,
    "tables": 4,
    "figures": 5,
    "equations": 3,
    "callouts": 3,
    "notation_boxes": 3,
    "fills": 9,
    "list_items": 21,
}

PALETTE = {
    "page": "#f9f9f7",
    "surface": "#fcfcfb",
    "ink": "#0b0b0b",
    "ink-2": "#52514e",
    "ink-muted": "#898781",
    "rule": "#e1e0d9",
    "baseline": "#c3c2b7",
    "pos": "#2a78d6",
    "neg": "#e34948",
    "ser2": "#eb6834",
    "ser3": "#1baf7a",
    "accent": "#184f95",
}

WORD_COLORS = {
    "ink": "0B0B0B",
    "ink-2": "52514E",
    "ink-muted": "74726C",
    "rule": "E1E0D9",
    "surface": "F7F8FA",
    "accent": "184F95",
    "fill": "FFF2CC",
    "fill-edge": "B57D05",
    "table-head": "EEF3F8",
}

TABLE_PROPORTIONS = (
    (0.31, 0.39, 0.30),
    (0.18, 0.32, 0.50),
    (0.20, 0.40, 0.40),
    (0.28, 0.72),
)

SVG_RE = re.compile(r"<svg\b.*?</svg>", re.IGNORECASE | re.DOTALL)
FIXED_TIME = dt.datetime(2000, 1, 1, tzinfo=dt.timezone.utc)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a Google-Docs-friendly DOCX from docs/index.html."
    )
    parser.add_argument(
        "source",
        nargs="?",
        default="docs/index.html",
        type=Path,
        help="Controlled HTML source (default: docs/index.html)",
    )
    parser.add_argument(
        "output",
        nargs="?",
        default="deliverables/search-routing-quality-handoff.docx",
        type=Path,
        help="DOCX output path",
    )
    parser.add_argument(
        "--keep-figures",
        type=Path,
        help="Also write the fixed-light PNG figures to this directory.",
    )
    return parser.parse_args()


def class_tokens(element: etree._Element) -> set[str]:
    return set((element.get("class") or "").split())


def normalized_text(element: etree._Element) -> str:
    return " ".join("".join(element.itertext()).split())


def validate_source(tree: etree._Element, source_text: str) -> None:
    actual = {
        "h1": len(tree.xpath("//h1")),
        "h2": len(tree.xpath("//h2")),
        "h3": len(tree.xpath("//h3")),
        "tables": len(tree.xpath("//table")),
        "figures": len(tree.xpath("//figure")),
        "equations": len(tree.xpath('//*[contains(concat(" ", normalize-space(@class), " "), " eq ")]')),
        "callouts": len(tree.xpath('//*[contains(concat(" ", normalize-space(@class), " "), " callout ")]')),
        "notation_boxes": len(tree.xpath('//*[contains(concat(" ", normalize-space(@class), " "), " notn ")]')),
        "fills": len(tree.xpath('//*[contains(concat(" ", normalize-space(@class), " "), " fill ")]')),
        "list_items": len(tree.xpath("//li")),
    }
    minimum_counts = {"h3", "list_items"}
    errors = []
    for name in EXPECTED:
        if name == "fills":
            continue
        if name in minimum_counts and actual[name] < EXPECTED[name]:
            errors.append(
                f"{name}: expected at least {EXPECTED[name]}, found {actual[name]}"
            )
        elif name not in minimum_counts and actual[name] != EXPECTED[name]:
            errors.append(
                f"{name}: expected {EXPECTED[name]}, found {actual[name]}"
            )
    if not 0 <= actual["fills"] <= EXPECTED["fills"]:
        errors.append(
            f"fills: expected between 0 and {EXPECTED['fills']}, "
            f"found {actual['fills']}"
        )
    if len(SVG_RE.findall(source_text)) != EXPECTED["figures"]:
        errors.append(
            f"inline SVGs: expected {EXPECTED['figures']}, "
            f"found {len(SVG_RE.findall(source_text))}"
        )
    if errors:
        raise ValueError("Source invariant failure:\n  - " + "\n  - ".join(errors))


def figure_metadata(svg_text: str) -> tuple[str, str]:
    title_match = re.search(r"<title(?:\s[^>]*)?>(.*?)</title>", svg_text, re.I | re.S)
    desc_match = re.search(r"<desc(?:\s[^>]*)?>(.*?)</desc>", svg_text, re.I | re.S)
    title = html_stdlib.unescape(re.sub(r"<[^>]+>", "", title_match.group(1))).strip() if title_match else ""
    desc = html_stdlib.unescape(re.sub(r"<[^>]+>", "", desc_match.group(1))).strip() if desc_match else ""
    return title, " ".join(desc.split())


def fixed_light_svg(svg_text: str) -> str:
    rendered = svg_text
    for name, color in PALETTE.items():
        rendered = rendered.replace(f"var(--{name})", color)
    rendered = re.sub(
        r"<svg\b",
        '<svg xmlns="http://www.w3.org/2000/svg"',
        rendered,
        count=1,
        flags=re.I,
    )
    # Tooltip hit target and hidden hover markers are interaction-only.
    rendered = re.sub(
        r"<(?:rect|line|circle)\b[^>]*\bid=\"(?:lc-hit|lc-x|lc-p1|lc-p2)\"[^>]*/?>",
        "",
        rendered,
        flags=re.I | re.S,
    )
    return rendered


def render_figures(
    source_text: str, output_dir: Path
) -> tuple[list[Path], list[tuple[str, str]], list[str]]:
    converter = shutil.which("rsvg-convert")
    if not converter:
        raise RuntimeError(
            "rsvg-convert is required. Install librsvg "
            "(macOS: `brew install librsvg`; Debian/Ubuntu: "
            "`apt-get install librsvg2-bin`)."
        )
    output_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    metadata: list[tuple[str, str]] = []
    hashes: list[str] = []
    for index, raw_svg in enumerate(SVG_RE.findall(source_text), start=1):
        svg = fixed_light_svg(raw_svg)
        svg_path = output_dir / f"figure-{index:02d}.svg"
        png_path = output_dir / f"figure-{index:02d}.png"
        svg_path.write_text(svg, encoding="utf-8")
        subprocess.run(
            [
                converter,
                "--background-color",
                PALETTE["surface"],
                "--width",
                "2040",
                "--keep-aspect-ratio",
                "--output",
                str(png_path),
                str(svg_path),
            ],
            check=True,
        )
        paths.append(png_path)
        metadata.append(figure_metadata(raw_svg))
        hashes.append(hashlib.sha256(raw_svg.encode("utf-8")).hexdigest())
    return paths, metadata, hashes


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed(table, widths_twips: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(sum(widths_twips)))
    width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_twips in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width_twips))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_twips[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_keep(paragraph, *, next_: bool = False, together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = next_
    paragraph.paragraph_format.keep_together = together


def set_paragraph_bottom_border(paragraph, color: str = "E1E0D9", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def create_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink"])
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink"])
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    heading1 = styles["Heading 1"]
    heading1.font.name = "Arial"
    heading1.font.size = Pt(17)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink"])
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(9)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True

    heading2 = styles["Heading 2"]
    heading2.font.name = "Arial"
    heading2.font.size = Pt(12.5)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink"])
    heading2.paragraph_format.space_before = Pt(15)
    heading2.paragraph_format.space_after = Pt(5)
    heading2.paragraph_format.keep_with_next = True
    heading2.paragraph_format.keep_together = True

    for style_name in ("List Number", "List Bullet"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(6)

    if "Section Kicker" not in styles:
        kicker = styles.add_style("Section Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = styles["Section Kicker"]
    kicker.font.name = "Arial"
    kicker.font.size = Pt(8)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string(WORD_COLORS["accent"])
    kicker.paragraph_format.space_after = Pt(3)
    kicker.paragraph_format.keep_with_next = True

    if "Standfirst" not in styles:
        standfirst = styles.add_style("Standfirst", WD_STYLE_TYPE.PARAGRAPH)
    else:
        standfirst = styles["Standfirst"]
    standfirst.font.name = "Arial"
    standfirst.font.size = Pt(15)
    standfirst.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink-2"])
    standfirst.paragraph_format.space_after = Pt(20)
    standfirst.paragraph_format.line_spacing = 1.12

    if "Caption" in styles:
        caption = styles["Caption"]
    else:
        caption = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink-2"])
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.05
    caption.paragraph_format.keep_together = True

    if "Legend" not in styles:
        legend = styles.add_style("Legend", WD_STYLE_TYPE.PARAGRAPH)
    else:
        legend = styles["Legend"]
    legend.font.name = "Arial"
    legend.font.size = Pt(8.5)
    legend.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink-2"])
    legend.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legend.paragraph_format.space_before = Pt(4)
    legend.paragraph_format.space_after = Pt(1)
    legend.paragraph_format.keep_with_next = True

    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    equation.font.name = "Times New Roman"
    equation.font.size = Pt(12)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(8)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.keep_together = True


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "SEARCH ROUTING QUALITY  ·  PROBLEM FRAMING"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        set_paragraph_bottom_border(paragraph, WORD_COLORS["rule"], "6")
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(7.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink-muted"])

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_before = Pt(2)
        run = paragraph.add_run("Compared to what?   ·   ")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink-muted"])
        add_page_field(paragraph)


def add_text_run(
    paragraph,
    text: str | None,
    *,
    bold: bool = False,
    italic: bool = False,
    math: bool = False,
    fill: bool = False,
    subscript: bool = False,
    superscript: bool = False,
) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.subscript = subscript
    run.font.superscript = superscript
    if math:
        run.font.name = "Times New Roman"
        run.italic = True
    if fill:
        run.text = f"[FILL: {' '.join(text.split())}]"
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(WORD_COLORS["fill-edge"])
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def append_inline(
    element: etree._Element,
    paragraph,
    *,
    bold: bool = False,
    italic: bool = False,
    math: bool = False,
) -> None:
    add_text_run(paragraph, element.text, bold=bold, italic=italic, math=math)
    for child in element:
        tag = etree.QName(child).localname.lower()
        classes = class_tokens(child)
        if tag == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif "fill" in classes:
            add_text_run(paragraph, normalized_text(child), fill=True)
        elif tag in {"strong", "b"}:
            append_inline(child, paragraph, bold=True, italic=italic, math=math)
        elif tag in {"em", "i"}:
            append_inline(child, paragraph, bold=bold, italic=True, math=math)
        elif tag == "sub":
            add_text_run(
                paragraph,
                normalized_text(child),
                bold=bold,
                italic=italic,
                math=math,
                subscript=True,
            )
        elif tag == "sup":
            add_text_run(
                paragraph,
                normalized_text(child),
                bold=bold,
                italic=italic,
                math=math,
                superscript=True,
            )
        elif tag == "a":
            # The generic source currently has no links. Keeping the visible text
            # is safer than creating an unreviewed external relationship.
            append_inline(child, paragraph, bold=bold, italic=italic, math=math)
        else:
            append_inline(
                child,
                paragraph,
                bold=bold,
                italic=italic,
                math=math or "m" in classes,
            )
        add_text_run(paragraph, child.tail, bold=bold, italic=italic, math=math)


def add_body_paragraph(container, element: etree._Element, style: str | None = None):
    paragraph = container.add_paragraph(style=style)
    append_inline(element, paragraph)
    set_keep(paragraph, together=True)
    return paragraph


def add_equation(container, element: etree._Element) -> None:
    paragraph = container.add_paragraph(style="Equation")
    add_text_run(paragraph, element.text)
    annotation = None
    for child in element:
        if "under" in class_tokens(child):
            annotation = normalized_text(child)
        else:
            tag = etree.QName(child).localname.lower()
            append_inline(
                child,
                paragraph,
                italic=tag in {"i", "em"},
                math=tag in {"i", "em"},
            )
        add_text_run(paragraph, child.tail)
    if annotation:
        under = container.add_paragraph()
        under.alignment = WD_ALIGN_PARAGRAPH.CENTER
        under.paragraph_format.space_before = Pt(0)
        under.paragraph_format.space_after = Pt(7)
        run = under.add_run(annotation)
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink-muted"])
        set_keep(under, together=True)


def add_box(document: Document, element: etree._Element, *, accent: bool) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_table_fixed(table, [9360])
    set_cell_shading(cell, "F5F7F9" if accent else "F8F8F6")
    set_cell_margins(cell, top=180, start=240, bottom=150, end=240)
    border = {
        "val": "single",
        "sz": "14" if accent else "6",
        "color": WORD_COLORS["accent"] if accent else "D7D9DC",
    }
    thin = {"val": "single", "sz": "6", "color": "D7D9DC"}
    set_cell_borders(
        cell,
        start=border,
        top=thin,
        bottom=thin,
        end=thin,
        insideH={"val": "nil"},
        insideV={"val": "nil"},
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    for child in element:
        classes = class_tokens(child)
        tag = etree.QName(child).localname.lower()
        if "label" in classes:
            label = cell.add_paragraph()
            label.paragraph_format.space_after = Pt(5)
            label.paragraph_format.keep_with_next = True
            run = label.add_run(normalized_text(child).upper())
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WORD_COLORS["accent"])
        elif "eq" in classes:
            add_equation(cell, child)
        elif tag == "p":
            paragraph = add_body_paragraph(cell, child)
            paragraph.paragraph_format.space_after = Pt(5)
    if not cell.paragraphs:
        cell.add_paragraph()
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    spacer.paragraph_format.line_spacing = Pt(1)


def add_table(document: Document, element: etree._Element, index: int) -> None:
    rows = element.xpath("./thead/tr | ./tbody/tr | ./tr")
    if not rows:
        return
    columns = max(len(row.xpath("./th | ./td")) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = [round(9360 * proportion) for proportion in TABLE_PROPORTIONS[index]]
    widths[-1] += 9360 - sum(widths)
    set_table_fixed(table, widths)

    for row_index, source_row in enumerate(rows):
        target_row = table.rows[row_index]
        source_cells = source_row.xpath("./th | ./td")
        for col_index, source_cell in enumerate(source_cells):
            cell = target_row.cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell, top=120, start=130, bottom=120, end=130)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.03
            append_inline(source_cell, paragraph)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9.2)
                if row_index == 0 or col_index == 0:
                    run.bold = True
            if row_index == 0:
                set_cell_shading(cell, WORD_COLORS["table-head"])
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor.from_string(WORD_COLORS["ink-2"])
        target_row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def legend_color(span: etree._Element) -> str:
    style = " ".join(span.xpath(".//i/@style"))
    match = re.search(r"var\(--([^)]+)\)", style)
    if match and match.group(1) in PALETTE:
        return PALETTE[match.group(1)].lstrip("#").upper()
    return WORD_COLORS["ink-2"]


def add_figure(
    document: Document,
    figure: etree._Element,
    png_path: Path,
    metadata: tuple[str, str],
) -> None:
    title, description = metadata
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(7)
    image_paragraph.paragraph_format.space_after = Pt(2)
    set_keep(image_paragraph, next_=True, together=True)
    run = image_paragraph.add_run()
    inline_shape = run.add_picture(str(png_path), width=Inches(6.5))
    inline_shape._inline.docPr.set("descr", description or title)
    inline_shape._inline.docPr.set("title", title)

    legends = figure.xpath(
        './/*[contains(concat(" ", normalize-space(@class), " "), " legend ")]/span'
    )
    if legends:
        legend = document.add_paragraph(style="Legend")
        for index, span in enumerate(legends):
            if index:
                legend.add_run("    ")
            square = legend.add_run("■ ")
            square.font.color.rgb = RGBColor.from_string(legend_color(span))
            label = "".join(span.xpath(".//text()[not(ancestor::i)]")).strip()
            legend.add_run(label)
        set_keep(legend, next_=True, together=True)

    captions = figure.xpath("./figcaption")
    if captions:
        caption = document.add_paragraph(style="Caption")
        append_inline(captions[0], caption)
        set_keep(caption, together=True)


def add_title_block(document: Document, header: etree._Element) -> None:
    eyebrow = header.xpath(
        './/*[contains(concat(" ", normalize-space(@class), " "), " eyebrow ")]'
    )[0]
    paragraph = document.add_paragraph(style="Section Kicker")
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.add_run(normalized_text(eyebrow).upper())

    title = header.xpath("./h1")[0]
    title_paragraph = document.add_paragraph(style="Title")
    append_inline(title, title_paragraph)

    standfirst = header.xpath(
        './/*[contains(concat(" ", normalize-space(@class), " "), " standfirst ")]'
    )[0]
    standfirst_paragraph = document.add_paragraph(style="Standfirst")
    append_inline(standfirst, standfirst_paragraph)

    byline = header.xpath(
        './/*[contains(concat(" ", normalize-space(@class), " "), " byline ")]'
    )[0]
    byline_paragraph = document.add_paragraph()
    byline_paragraph.paragraph_format.space_before = Pt(4)
    byline_paragraph.paragraph_format.space_after = Pt(18)
    set_paragraph_bottom_border(byline_paragraph)
    append_inline(byline, byline_paragraph)


def add_section_heading(document: Document, element: etree._Element) -> None:
    number_spans = [
        child for child in element if "num" in class_tokens(child)
    ]
    kicker_text = normalized_text(number_spans[0]) if number_spans else ""
    full_text = normalized_text(element)
    title_text = full_text[len(kicker_text) :].strip() if kicker_text else full_text

    kicker = document.add_paragraph(style="Section Kicker")
    kicker.paragraph_format.page_break_before = True
    kicker.paragraph_format.keep_with_next = True
    kicker.add_run(kicker_text.upper())

    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(title_text)


def add_footer_content(document: Document, element: etree._Element) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(20)
    set_paragraph_bottom_border(paragraph)
    append_inline(element, paragraph)


def build_document(
    tree: etree._Element,
    figures: list[Path],
    figure_metadata_items: list[tuple[str, str]],
) -> Document:
    document = Document()
    create_styles(document)
    configure_sections(document)
    document.core_properties.title = "Compared to what? — Search routing quality"
    document.core_properties.subject = (
        "Problem framing and analytical methodology for search routing quality"
    )
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.created = FIXED_TIME
    document.core_properties.modified = FIXED_TIME
    document.core_properties.revision = 1
    document.core_properties.keywords = ""
    document.core_properties.comments = ""

    wrap = tree.xpath(
        '//*[contains(concat(" ", normalize-space(@class), " "), " wrap ")]'
    )[0]
    table_index = 0
    figure_index = 0
    for child in wrap:
        tag = etree.QName(child).localname.lower()
        classes = class_tokens(child)
        if tag == "header":
            add_title_block(document, child)
        elif tag == "h2":
            add_section_heading(document, child)
        elif tag == "h3":
            paragraph = document.add_paragraph(style="Heading 2")
            append_inline(child, paragraph)
        elif tag == "p":
            add_body_paragraph(document, child)
        elif tag in {"ol", "ul"}:
            style = "List Number" if tag == "ol" else "List Bullet"
            for item in child.xpath("./li"):
                paragraph = document.add_paragraph(style=style)
                append_inline(item, paragraph)
                set_keep(paragraph, together=True)
        elif tag == "table":
            add_table(document, child, table_index)
            table_index += 1
        elif tag == "figure":
            add_figure(
                document,
                child,
                figures[figure_index],
                figure_metadata_items[figure_index],
            )
            figure_index += 1
        elif "callout" in classes:
            add_box(document, child, accent=True)
        elif "notn" in classes:
            add_box(document, child, accent=False)
        elif tag == "footer":
            add_footer_content(document, child)
        elif tag == "hr":
            continue
    return document


def normalize_docx(source: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    core_ns = {
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
        "dc": "http://purl.org/dc/elements/1.1/",
        "dcterms": "http://purl.org/dc/terms/",
    }
    with zipfile.ZipFile(source) as archive:
        members = {
            name: archive.read(name)
            for name in archive.namelist()
            if name != "docProps/custom.xml"
            and name != "docProps/thumbnail.jpeg"
            and not name.startswith("customXml/")
        }
    core_name = "docProps/core.xml"
    if core_name in members:
        root = etree.fromstring(members[core_name])
        for xpath in (
            "dc:creator",
            "cp:lastModifiedBy",
            "cp:keywords",
            "dc:description",
        ):
            for node in root.xpath(xpath, namespaces=core_ns):
                node.text = ""
        for xpath in ("dcterms:created", "dcterms:modified"):
            for node in root.xpath(xpath, namespaces=core_ns):
                node.text = "2000-01-01T00:00:00Z"
        members[core_name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    app_name = "docProps/app.xml"
    if app_name in members:
        root = etree.fromstring(members[app_name])
        app_ns = {
            "ep": (
                "http://schemas.openxmlformats.org/officeDocument/"
                "2006/extended-properties"
            )
        }
        for node in root.xpath("ep:Template", namespaces=app_ns):
            node.text = ""
        for node in root.xpath("ep:Application", namespaces=app_ns):
            node.text = "Search Routing Quality DOCX builder"
        for node in root.xpath("ep:AppVersion", namespaces=app_ns):
            node.text = "1.0"
        members[app_name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    relationship_files = (
        "_rels/.rels",
        "word/_rels/document.xml.rels",
    )
    for name in relationship_files:
        if name not in members:
            continue
        root = etree.fromstring(members[name])
        for relationship in list(root):
            rel_type = relationship.get("Type", "")
            if rel_type.endswith("/customXml") or rel_type.endswith(
                "/metadata/thumbnail"
            ):
                root.remove(relationship)
        members[name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    content_types_name = "[Content_Types].xml"
    if content_types_name in members:
        root = etree.fromstring(members[content_types_name])
        for node in list(root):
            part_name = node.get("PartName", "")
            if part_name.startswith("/customXml/"):
                root.remove(node)
        members[content_types_name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name in sorted(members):
            info = zipfile.ZipInfo(name, (2000, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            archive.writestr(info, members[name])


def main() -> int:
    args = parse_args()
    source = args.source.resolve()
    output = args.output.resolve()
    source_text = source.read_text(encoding="utf-8")
    tree = html.fromstring(source_text)
    validate_source(tree, source_text)

    with tempfile.TemporaryDirectory(prefix="search-routing-docx-") as tmp:
        tmp_dir = Path(tmp)
        figures, metadata, hashes = render_figures(source_text, tmp_dir / "figures")
        document = build_document(tree, figures, metadata)
        unsanitized = tmp_dir / "handoff-unsanitized.docx"
        document.save(unsanitized)
        normalize_docx(unsanitized, output)

        if args.keep_figures:
            figure_dir = args.keep_figures.resolve()
            figure_dir.mkdir(parents=True, exist_ok=True)
            for index, path in enumerate(figures, start=1):
                shutil.copy2(path, figure_dir / f"figure-{index:02d}.png")
            manifest = figure_dir / "SOURCE_SVG_SHA256.txt"
            manifest.write_text(
                "\n".join(
                    f"figure-{index:02d}.svg  {digest}"
                    for index, digest in enumerate(hashes, start=1)
                )
                + "\n",
                encoding="utf-8",
            )

    print(f"Built {output}")
    print(f"SHA256 {hashlib.sha256(output.read_bytes()).hexdigest()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
